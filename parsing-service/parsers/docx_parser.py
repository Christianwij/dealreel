from docx import Document
from loguru import logger
import io
import time
from typing import List, Dict, Any

from exceptions import (
    DocumentCorruptedError,
    TextExtractionError,
    TableExtractionError,
    MetadataExtractionError
)
from monitoring import record_extraction_metric, record_extraction_time

async def parse_docx(file_content: bytes) -> dict:
    """Parse DOCX file content and return structured data."""
    try:
        # Create document from bytes
        doc_stream = io.BytesIO(file_content)
        doc = Document(doc_stream)
        
        # Extract metadata
        metadata_start = time.time()
        try:
            core_properties = doc.core_properties
            metadata = {
                "title": core_properties.title or "",
                "author": core_properties.author or "",
                "subject": core_properties.subject or "",
                "keywords": core_properties.keywords or "",
                "created": core_properties.created.isoformat() if core_properties.created else "",
                "modified": core_properties.modified.isoformat() if core_properties.modified else "",
            }
            record_extraction_time("metadata", "docx", time.time() - metadata_start)
        except Exception as e:
            logger.error(f"Failed to extract DOCX metadata: {str(e)}")
            raise MetadataExtractionError(
                "Failed to extract DOCX metadata",
                document_type="docx",
                details={"error": str(e)}
            )
        
        # Extract text and tables
        text_content = []
        tables = []
        
        text_start = time.time()
        try:
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
                    record_extraction_metric("paragraphs", "docx")
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                try:
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        if any(row_data):  # Skip empty rows
                            table_data.append(row_data)
                    if table_data:
                        tables.append(table_data)
                        record_extraction_metric("tables", "docx")
                except Exception as e:
                    logger.warning(f"Failed to extract table: {str(e)}")
                    continue
            
            record_extraction_time("text", "docx", time.time() - text_start)
            
        except Exception as e:
            logger.error(f"Failed to extract DOCX content: {str(e)}")
            raise TextExtractionError(
                "Failed to extract DOCX content",
                document_type="docx",
                details={"error": str(e)}
            )
        
        return {
            "content_type": "docx",
            "metadata": metadata,
            "text": "\n".join(text_content),
            "tables": tables,
            "sections": len(doc.sections)
        }
        
    except (TextExtractionError, MetadataExtractionError) as e:
        raise e
    except Exception as e:
        logger.error(f"Failed to parse DOCX: {str(e)}")
        raise DocumentCorruptedError(
            "Failed to parse DOCX document",
            document_type="docx",
            details={"error": str(e)}
        )
    finally:
        if 'doc_stream' in locals():
            doc_stream.close() 